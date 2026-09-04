"""Gera DOCX do Anexo STAR Ondas de Calor no formato solicitado."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "apresentacoes" / "STAR_Anexo1_Ondas_de_Calor_MT.docx"


def _set_run(run, *, bold=False, size=11, color=None):
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run(r, bold=True, size=16, color=RGBColor(0x1D, 0x35, 0x7F))


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        _set_run(run, bold=True, size=13 if level == 1 else 12, color=RGBColor(0x1D, 0x35, 0x7F))


def add_para(doc, text, *, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run(r, bold=bold, size=11)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_run(r, size=11)
    p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        _set_run(r, bold=True, size=10)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            _set_run(r, size=10)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title(doc, "Oficina STAR")
    add_title(doc, "MATERIAIS NECESSÁRIOS")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tema: Ondas de Calor")
    _set_run(r, bold=True, size=13)

    add_para(doc, "Estado: Mato Grosso")
    add_para(doc, "Órgão: CIEVS-MT / Secretaria de Estado de Saúde de Mato Grosso")
    add_para(
        doc,
        "Referência de dados: SE 35/2026 (atualização 03/09/2026). Itens sem série completa "
        "de 5 anos ou sem dado disponível são registrados como parciais ou indisponíveis.",
    )

    # ---- SEÇÃO 1 ----
    add_heading(doc, "1. Análise epidemiológica e climática (últimos 5 anos, no mínimo)", 1)

    add_heading(
        doc,
        "1.1 Série histórica de ondas de calor, temperaturas máximas, anomalias de temperatura "
        "e alertas meteorológicos, por ano, mês e/ou semana epidemiológica, no estado",
        2,
    )
    add_para(doc, "Situação: parcial.", bold=True)
    add_bullet(doc, "Histórico climático municipal acumulado: 28/04/2026 a 09/09/2026 (abaixo do mínimo de 5 anos).")
    add_bullet(doc, "SE 35/2026: 115/142 (81,0%) municípios em vermelho ou roxo; projeção ~7 dias: 128/142 (90,1%).")
    add_bullet(doc, "Tmáx máxima na rodada: 39,3 °C; 15 municípios com Tmáx ≥ 37 °C.")
    add_bullet(doc, "Pico recente em Cuiabá: 40,4 °C em 31/08/2026.")
    add_bullet(doc, "Alertas INMET recentes: predominância de baixa umidade; sem arquivo histórico completo de 5 anos.")
    add_bullet(doc, "Lacuna principal: série climatológica municipal ≥5 anos de ondas, anomalias e alertas por ano/mês/SE.")

    add_heading(
        doc,
        "1.2 Definição dos limiares utilizados para caracterizar onda de calor no território, "
        "incluindo duração, intensidade e áreas abrangidas",
        2,
    )
    add_para(doc, "Situação: disponível (definição operacional estadual).", bold=True)
    add_table(
        doc,
        ["Critério", "Definição"],
        [
            ("UTCI/proxy", "Verde ≤26; Amarela ≤32; Laranja ≤38; Vermelha ≤46; Roxa acima ou por combinação"),
            ("Tmáx (fallback)", "Amarela ≥37 °C; Laranja ≥39 °C; Vermelha ≥41 °C; Roxa ≥43 °C"),
            ("Risco cumulativo 3 dias", "Umbral 39 °C; Amarela ≥3; Laranja ≥7; Vermelha ≥12; Roxa ≥18"),
            ("EHF adaptado", "Positivo >0; persistência de emergência: 5 dias"),
            ("Onda de calor (duração)", "≥ 2 dias consecutivos com temperatura média acima do P95 local"),
            ("Abrangência", "Estadual, com classificação por município e região de saúde"),
        ],
    )
    add_para(
        doc,
        "Nota: o P95 municipal, na ausência de climatologia longa, pode operar como proxy relativo "
        "e não deve ser interpretado como anomalia climatológica oficial 1981–2010.",
    )

    add_heading(
        doc,
        "1.3 Mapas temáticos com municípios/regiões mais expostos ou afetados, áreas urbanas com "
        "ilhas de calor, baixa arborização, alta densidade populacional ou maior vulnerabilidade social",
        2,
    )
    add_para(doc, "Situação: parcial.", bold=True)
    add_bullet(doc, "Disponível: mapas de classificação atual e projeção ~7 dias; territórios tradicionais; densidade, idosos, crianças, ruralidade e índice de vulnerabilidade ao calor.")
    add_bullet(doc, "Mais expostos termicamente (Tmáx): Cocalinho (39,3 °C), Nova Nazaré (38,7 °C), Araguaiana (38,6 °C), Novo Santo Antônio (38,6 °C), Canabrava do Norte (38,0 °C).")
    add_bullet(doc, "Maior vulnerabilidade ao calor (exemplos): São José do Povo, Nossa Senhora do Livramento, Porto Estrela, Vale de São Domingos, Acorizal.")
    add_bullet(doc, "Indisponível: mapas de ilhas de calor urbanas, baixa arborização e NDVI/cobertura vegetal intraurbana.")

    add_heading(
        doc,
        "1.4 Identificação de tendências, sazonalidade, períodos de maior risco e previsão climática/sazonal, quando disponível",
        2,
    )
    add_para(doc, "Situação: parcial.", bold=True)
    add_bullet(doc, "Período crítico operacional: julho a novembro.")
    add_bullet(doc, "Contexto: El Niño 2026–2027; temperatura acima da média; estiagem/transição chuvosa irregular.")
    add_bullet(doc, "Predição operacional de curto prazo (~7 dias) disponível por município.")
    add_bullet(doc, "Chuva em 01/09/2026 reduziu temporariamente municípios com Tmáx ≥ 37 °C, sem encerrar o risco alto estadual.")
    add_bullet(doc, "Previsão sazonal oficial (CPTEC/INMET/FUNCEME) usada como referência qualitativa; sem série própria de 5 anos neste material.")

    add_heading(
        doc,
        "1.5 Dados de atendimentos, internações, remoções e óbitos possivelmente relacionados ao calor "
        "(desidratação, insolação, exaustão pelo calor e agravamento de doenças cardiovasculares, respiratórias e renais)",
        2,
    )
    add_para(doc, "Situação: parcial.", bold=True)
    add_bullet(doc, "CIDs/grupos monitorados: E86, E87, T67, X30 e agravos cardiocirculatórios, respiratórios e renais sensíveis a calor.")
    add_bullet(doc, "Óbitos sensíveis a calor (SIM agregado): 15.329 registros (aprox. jan/2024 a ago/2026).")
    add_bullet(doc, "Lacunas: atendimentos ambulatoriais específicos e remoções não rotinizados; tipificação de calor direto frágil; janela <5 anos completa.")

    add_heading(
        doc,
        "1.6 Distribuição dos eventos/agravos por faixa etária, sexo, município/região de saúde, "
        "local de ocorrência e grupos populacionais vulneráveis",
        2,
    )
    add_para(doc, "Situação: parcial.", bold=True)
    add_bullet(doc, "Disponível: distribuição por município e região de saúde.")
    add_bullet(doc, "Sexo/idade/faixa etária existem no SIM bruto, mas não consolidados em tabela rotineira específica de ondas de calor neste material.")
    add_bullet(doc, "Grupos com dado municipal: idosos, crianças 0–4 anos, população rural e territórios tradicionais.")
    add_bullet(doc, "Sem distribuição rotineira para gestantes, imunossuprimidos, situação de rua e população privada de liberdade.")

    add_heading(
        doc,
        "1.7 Informações sobre aumento de demanda ou sobrecarga dos serviços de saúde em períodos "
        "de calor extremo/ondas de calor",
        2,
    )
    add_para(doc, "Situação: disponível como proxy operacional.", bold=True)
    add_bullet(doc, "Ocupação hospitalar IndicaSUS (03/09/2026): 57,0% (3.346/5.872 leitos), em 85 municípios com ocupação informada.")
    add_bullet(doc, "Pressão assistencial (SISREG/índice de pressão) monitorada por município, distinta da ocupação de leitos.")

    add_heading(doc, "1.8 Outras análises que a área técnica julgar pertinentes", 2)
    add_bullet(doc, "Sobreposição calor e qualidade do ar: 18/142 municípios com PM2,5 ≥ 25 µg/m³.")
    add_bullet(doc, "Persistência de focos de calor/queimadas no período.")
    add_bullet(doc, "Alívio térmico temporário após 01/09/2026, com retorno do risco na projeção de curto prazo.")

    # ---- SEÇÃO 2 ----
    add_heading(doc, "2. Caracterização do evento no território e seus impactos à saúde", 1)

    add_heading(
        doc,
        "2.1 Principais impactos diretos à saúde relacionados ao calor já observados no estado, com dados, quando disponíveis",
        2,
    )
    add_bullet(doc, "Monitoramento de desidratação, insolação/exaustão pelo calor e agravamento cardiovascular, respiratório e renal.")
    add_bullet(doc, "SIM — óbitos sensíveis a calor: 15.329 (jan/2024–ago/2026).")
    add_bullet(doc, "SE 35/2026: 81% dos municípios em vermelho/roxo eleva o risco de agravamentos agudos em grupos sensíveis.")

    add_heading(
        doc,
        "2.2 Possíveis impactos indiretos já identificados no estado (insegurança hídrica, qualidade da água e dos alimentos, "
        "aumento de DTHA, agravamento da qualidade do ar e interrupção de atividades essenciais), com dados, quando disponíveis",
        2,
    )
    add_table(
        doc,
        ["Impacto indireto", "Situação / dado"],
        [
            ("Insegurança hídrica / estiagem", "Parcial — monitoramento hidrológico em recortes"),
            ("Qualidade da água e dos alimentos", "Indisponível de forma consolidada neste material"),
            ("Aumento de DTHA", "Parcial — vigilância existe; sem série causal calor–DTHA de 5 anos aqui"),
            ("Agravamento da qualidade do ar", "Disponível — 18 municípios com PM2,5 ≥ 25 µg/m³"),
            ("Interrupção de atividades essenciais", "Indisponível de forma padronizada neste material"),
        ],
    )

    add_heading(
        doc,
        "2.3 Fatores que podem ampliar o risco (baixa umidade, dificuldade de acesso à água, moradias sem ventilação adequada, "
        "trabalho ao ar livre, eventos de massa, abrigos, alojamentos e áreas de grande circulação de pessoas)",
        2,
    )
    add_bullet(doc, "Baixa umidade relativa (avisos INMET frequentes no período recente).")
    add_bullet(doc, "Dificuldade de acesso à água em contextos de estiagem (parcialmente monitorada).")
    add_bullet(doc, "Trabalho ao ar livre / exposição ocupacional.")
    add_bullet(doc, "Alta densidade e circulação em áreas urbanas (ex.: Cuiabá e entorno).")
    add_bullet(doc, "Sobreposição calor + fumaça.")
    add_bullet(doc, "Moradias sem ventilação adequada, eventos de massa, abrigos e alojamentos: reconhecidos como fatores de risco, sem indicador municipal rotineiro consolidado neste material.")

    add_heading(
        doc,
        "2.4 Grupos com maior vulnerabilidade no estado (crianças, idosos, gestantes, pessoas com doenças crônicas, "
        "imunossuprimidos, pessoas em situação de rua, população privada de liberdade, população indígena, "
        "trabalhadores expostos ao calor e pessoas em moradias precárias), com dados, quando disponíveis",
        2,
    )
    add_table(
        doc,
        ["Grupo", "Dado disponível?", "Observação"],
        [
            ("Crianças", "Sim", "Proporção 0–4 anos por município"),
            ("Idosos", "Sim", "Proporção/contingente 60+ e índice de vulnerabilidade ao calor"),
            ("Gestantes", "Não", "Sem indicador municipal rotineiro neste material"),
            ("Doenças crônicas", "Não (individual)", "Necessita cruzamento assistencial pontual"),
            ("Imunossuprimidos", "Não", "Lacuna"),
            ("Situação de rua", "Não", "Lacuna — articulação com assistência social"),
            ("Privados de liberdade", "Não", "Lacuna"),
            ("População indígena", "Parcial", "Territórios tradicionais; articulação DSEI/SESAI"),
            ("Trabalhadores expostos", "Parcial", "Saúde do Trabalhador / AdaptaSUS"),
            ("Moradias precárias", "Parcial", "Proxy via ruralidade e índice de vulnerabilidade"),
        ],
    )

    add_heading(doc, "2.5 Outras características relevantes para a análise do evento no território", 2)
    add_bullet(doc, "Contexto de El Niño e emergência ambiental por queimadas.")
    add_bullet(doc, "Abrangência estadual, com maior concentração térmica recente no arco leste/nordeste e risco persistente na Baixada Cuiabana.")
    add_bullet(doc, "Risco multifatorial: calor + baixa umidade + fumaça + vulnerabilidade demográfica.")

    # ---- SEÇÃO 3 ----
    add_heading(doc, "3. Informações sobre o Sistema de Vigilância e Capacidade de Resposta", 1)

    add_heading(
        doc,
        "3.1 Estrutura atual da vigilância dos eventos associados a Ondas de Calor (níveis federal, estadual e municipal)",
        2,
    )
    add_table(
        doc,
        ["Nível", "Estrutura"],
        [
            ("Federal", "INMET; Ministério da Saúde / AdaptaSUS; Rede CIEVS"),
            ("Estadual", "CIEVS-MT; Sala de Situação (Portaria nº 0590/2026/GBSES); ARARAS MT; Vigidesastres; áreas técnicas da SES-MT"),
            ("Municipal", "Vigilância em Saúde; Defesa Civil; APS e rede assistencial; uso de alertas e painéis estaduais"),
        ],
    )

    add_heading(doc, "3.2 Disponibilidade de insumos, exames laboratoriais e protocolos", 2)
    add_bullet(doc, "Protocolos/limiares operacionais: disponíveis.")
    add_bullet(doc, "Insumos e exames laboratoriais: informação parcial; sem inventário quantitativo completo por município específico para ondas de calor neste material.")

    add_heading(doc, "3.3 Recursos humanos envolvidos e eventuais lacunas identificadas", 2)
    add_bullet(doc, "Envolvimento técnico do CIEVS-MT, Sala de Situação, áreas de vigilância e assistência da SES-MT e pontos focais municipais/regionais.")
    add_bullet(doc, "Lacuna: ausência de quadro formal consolidado de RH dedicado exclusivamente a ondas de calor neste levantamento.")

    add_heading(doc, "3.4 Articulação com outros setores (meio ambiente, educação, assistência social, etc.)", 2)
    add_bullet(doc, "Meio ambiente / SEMA e correlatos de emergência ambiental.")
    add_bullet(doc, "Corpo de Bombeiros; Educação; Assistência social; Defesa Civil.")
    add_bullet(doc, "Saúde do Trabalhador; DSEI/SESAI; demais áreas técnicas da SES no Plano El Niño / Sala de Situação.")

    add_heading(
        doc,
        "3.5 Existência de planos de contingência, políticas públicas, normas técnicas e legislação específica relacionadas à doença",
        2,
    )
    add_bullet(doc, "Portaria nº 0590/2026/GBSES — Sala de Situação em Saúde (El Niño 2026–2027 e extremos climáticos).")
    add_bullet(doc, "Plano El Niño SES-MT.")
    add_bullet(doc, "Decreto nº 2.015/2026 — emergência ambiental.")
    add_bullet(doc, "Diretrizes AdaptaSUS / mudanças climáticas e saúde.")
    add_bullet(doc, "ARARAS MT como apoio ao monitoramento integrado (complementar aos sistemas do SUS).")
    add_bullet(doc, "Planos municipais específicos de calor extremo: não catalogados de forma completa neste material.")

    add_heading(doc, "3.6 Principais desafios e oportunidades de melhoria do sistema de vigilância", 2)
    add_para(doc, "Desafios", bold=True)
    add_bullet(doc, "Consolidar série histórica ≥5 anos de ondas de calor, anomalias e alertas.")
    add_bullet(doc, "Produzir mapas de ilhas de calor e arborização.")
    add_bullet(doc, "Melhorar tipificação e estratificação (idade/sexo) de atendimentos, internações e óbitos relacionados ao calor.")
    add_bullet(doc, "Completar inventário de RH, insumos e planos municipais de contingência.")
    add_bullet(doc, "Integrar continuamente qualidade do ar, hídrica e sobrecarga assistencial ao indicador de onda de calor.")
    add_para(doc, "Oportunidades", bold=True)
    add_bullet(doc, "Rotina semanal padronizada via Sala de Situação e ARARAS MT.")
    add_bullet(doc, "Ampliar série climática e sanitária para 5 anos.")
    add_bullet(doc, "Fortalecer articulação intersetorial do Plano El Niño.")
    add_bullet(doc, "Priorizar municípios com alta exposição térmica e alta vulnerabilidade demográfica.")

    add_para(
        doc,
        "Documento elaborado para atendimento ao Anexo 1 — Avaliação STAR (Materiais_Ondas de Calor). CIEVS-MT / SES-MT.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
