"""Gera DOCX da minuta de Portaria ARARAS MT v3 e do encaminhamento."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
INST = ROOT / "docs" / "institucional"


def _set_run_font(run, *, bold: bool = False, italic: bool = False, size: float = 12) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")


def _add_mixed(paragraph, text: str, *, size: float = 12, center: bool = False) -> None:
    """Suporta **negrito** simples no texto."""
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, bold=True, size=size)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, size=size)


def _style_paragraph(paragraph, *, space_after: float = 6, first_line: bool = False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line:
        pf.first_line_indent = Cm(1.25)


def md_to_docx(md_path: Path, out_path: Path, *, skip_title_h1: bool = True) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    skip_until_blank_after_header = True
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Pular bloco de metadados inicial até o primeiro --- após o H1
        if skip_until_blank_after_header:
            if line.startswith("# ") and skip_title_h1:
                i += 1
                continue
            if line.startswith("**Arquivo:**") or line.startswith("**Relação") or line.startswith("**Finalidade") or line.startswith("**Data de") or line.startswith("**Uso:**") or line.startswith("**Anexos:**"):
                i += 1
                continue
            if line == "---":
                skip_until_blank_after_header = False
                i += 1
                continue
            if not line:
                i += 1
                continue

        if not line:
            i += 1
            continue

        if line == "---":
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            _style_paragraph(p, space_after=10)
            _add_mixed(p, line[3:].strip(), size=12, center=False)
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        if line.startswith("|") and "---" not in line:
            # Tabela markdown simples → parágrafo tabular
            cells = [c.strip() for c in line.strip("|").split("|")]
            p = doc.add_paragraph()
            _style_paragraph(p, space_after=4)
            _add_mixed(p, " | ".join(cells), size=10)
            i += 1
            continue

        if re.match(r"^\|?\s*-+\s*\|", line):
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            _style_paragraph(p, space_after=8)
            _add_mixed(p, line[2:], size=11)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Cabeçalho institucional / ementa / artigos
        is_header_block = line in {
            "GOVERNO DO ESTADO DE MATO GROSSO",
            "SECRETARIA DE ESTADO DE SAÚDE",
            "GABINETE DO SECRETÁRIO",
            "CENTRO DE INFORMAÇÕES ESTRATÉGICAS EM VIGILÂNCIA EM SAÚDE – CIEVS-MT",
            "REGISTRE-SE. PUBLIQUE-SE. CUMPRA-SE.",
        } or line.startswith("**PORTARIA Nº") or line.startswith("**DESPACHO")

        p = doc.add_paragraph()
        if is_header_block or line.startswith("**JULIANO") or line.startswith("Cuiabá"):
            _style_paragraph(p, space_after=4)
            _add_mixed(p, line, size=12, center=True)
        elif line.startswith("**Art.") or line.startswith("**§") or re.match(r"^[IVX]+\s*–", line) or line.startswith("I –") or line.startswith("II –") or line.startswith("III –") or line.startswith("IV –") or line.startswith("V –") or line.startswith("VI –") or line.startswith("VII –") or line.startswith("VIII –") or line.startswith("IX –") or line.startswith("X –") or line.startswith("XI –"):
            _style_paragraph(p, space_after=6, first_line=True)
            _add_mixed(p, line, size=12)
        elif line.startswith("CONSIDERANDO") or line.startswith("O SECRETÁRIO") or line.startswith("**RESOLVE:**") or line.startswith("À") or line.startswith("Senhor") or re.match(r"^\d+\.", line) or line.startswith("– ") or line.startswith("Atenciosamente"):
            _style_paragraph(p, space_after=6, first_line=line.startswith("CONSIDERANDO") or line.startswith("O SECRETÁRIO"))
            _add_mixed(p, line, size=12)
        else:
            _style_paragraph(p, space_after=6, first_line=False)
            _add_mixed(p, line, size=12)

        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[OK] {out_path}")


def main() -> None:
    md_to_docx(
        INST / "Minuta_Portaria_ARARAS_MT_v3.md",
        INST / "Minuta_Portaria_ARARAS_MT_v3.docx",
    )
    md_to_docx(
        INST / "Encaminhamento_Portaria_ARARAS_MT.md",
        INST / "Encaminhamento_Portaria_ARARAS_MT.docx",
        skip_title_h1=True,
    )


if __name__ == "__main__":
    main()
